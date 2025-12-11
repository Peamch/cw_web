#!/usr/bin/env python3
"""
–†–æ–∑—à–∏—Ä–µ–Ω–∏–π —Å–∫—Ä–∏–ø—Ç –¥–ª—è –≥–µ–Ω–µ—Ä–∞—Ü—ñ—ó –∫—É—Ä—Å–æ–≤–æ—ó —Ä–æ–±–æ—Ç–∏ –∑ –≤–µ–ª–∏–∫–æ—é –∫—ñ–ª—å–∫—ñ—Å—Ç—é –ø—Ä–∏–∫–ª–∞–¥—ñ–≤ –∫–æ–¥—É —Ç–∞ –ø–æ—è—Å–Ω–µ–Ω–Ω—è–º–∏.
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_font(run, name='Times New Roman', size=14, bold=False, color=None):
    """–í—Å—Ç–∞–Ω–æ–≤–ª—é—î —à—Ä–∏—Ñ—Ç –¥–ª—è —Ç–µ–∫—Å—Ç—É."""
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), name)

def add_para(doc, text, bold=False, color=None):
    """–î–æ–¥–∞—î –ø–∞—Ä–∞–≥—Ä–∞—Ñ –∑ —Ñ–æ—Ä–º–∞—Ç—É–≤–∞–Ω–Ω—è–º."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, bold=bold, color=color)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0.5)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    return p

def add_code_block(doc, code_lines, language="Java"):
    """–î–æ–¥–∞—î –±–ª–æ–∫ –∫–æ–¥—É –∑ —Å—ñ—Ä–∏–º —Ñ–æ–Ω–æ–º."""
    p = doc.add_paragraph()
    run = p.add_run('\n'.join(code_lines))
    set_font(run, name='Courier New', size=10)
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    # –î–æ–¥–∞—î–º–æ —Å—ñ—Ä–∏–π —Ñ–æ–Ω
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), 'F5F5F5')
    p._element.get_or_add_pPr().append(shading_elm)
    return p

def add_explanation(doc, title, text):
    """–î–æ–¥–∞—î –ø–æ—è—Å–Ω–µ–Ω–Ω—è –ø–µ—Ä–µ–¥ –∫–æ–¥–æ–º."""
    p = doc.add_paragraph()
    run = p.add_run(title)
    set_font(run, bold=True, size=13)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_before = Pt(12)
    
    p2 = doc.add_paragraph()
    run2 = p2.add_run(text)
    set_font(run2)
    p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p2.paragraph_format.first_line_indent = Inches(0.5)
    p2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    return p2

# –ó–∞–≤–∞–Ω—Ç–∞–∂—É—î–º–æ —ñ—Å–Ω—É—é—á–∏–π –¥–æ–∫—É–º–µ–Ω—Ç
doc = Document('/home/runner/work/cw_web/cw_web/–ö—É—Ä—Å–æ–≤–∞_–†–æ–±–æ—Ç–∞_HabitTracker.docx')

# –î–æ–¥–∞—î–º–æ –†–û–ó–î–Ü–õ 4.1 –∑ –ø—Ä–∏–∫–ª–∞–¥–∞–º–∏ –∫–æ–¥—É backend
doc.add_page_break()
doc.add_heading('–†–û–ó–î–Ü–õ 4. –†–ï–ê–õ–Ü–ó–ê–¶–Ü–Ø –Ü–ù–§–û–†–ú–ê–¶–Ü–ô–ù–û–á –°–ò–°–¢–ï–ú–ò', level=1)
add_para(doc, '–£ —Ü—å–æ–º—É —Ä–æ–∑–¥—ñ–ª—ñ –¥–µ—Ç–∞–ª—å–Ω–æ –æ–ø–∏—Å–∞–Ω–æ —Ä–µ–∞–ª—ñ–∑–∞—Ü—ñ—é –≤–µ–±-—Å–∏—Å—Ç–µ–º–∏ –¥–ª—è –≤—ñ–¥—Å—Ç–µ–∂–µ–Ω–Ω—è –∑–≤–∏—á–æ–∫ —Ç–∞ –¥–æ—Å—è–≥–Ω–µ–Ω—å. –ü—Ä–µ–¥—Å—Ç–∞–≤–ª–µ–Ω–æ –∞—Ä—Ö—ñ—Ç–µ–∫—Ç—É—Ä—É —Å–∏—Å—Ç–µ–º–∏, —Å—Ç—Ä—É–∫—Ç—É—Ä—É –ø—Ä–æ–µ–∫—Ç—É —Ç–∞ –∫–æ–Ω–∫—Ä–µ—Ç–Ω—ñ –ø—Ä–∏–∫–ª–∞–¥–∏ –∫–æ–¥—É backend —ñ frontend —á–∞—Å—Ç–∏–Ω –∑ –¥–µ—Ç–∞–ª—å–Ω–∏–º–∏ –ø–æ—è—Å–Ω–µ–Ω–Ω—è–º–∏.')

doc.add_heading('4.1 –†–µ–∞–ª—ñ–∑–∞—Ü—ñ—è backend —á–∞—Å—Ç–∏–Ω–∏ —Å–∏—Å—Ç–µ–º–∏', level=2)

# 4.1.1 –ú–æ–¥–µ–ª—å –¥–∞–Ω–∏—Ö - Entity –∫–ª–∞—Å–∏
doc.add_heading('4.1.1 –ú–æ–¥–µ–ª—å –¥–∞–Ω–∏—Ö - Entity –∫–ª–∞—Å–∏', level=3)

add_explanation(doc, 
    '–ö–ª–∞—Å User - –º–æ–¥–µ–ª—å –∫–æ—Ä–∏—Å—Ç—É–≤–∞—á–∞ —Å–∏—Å—Ç–µ–º–∏:',
    '–ö–ª–∞—Å User –ø—Ä–µ–¥—Å—Ç–∞–≤–ª—è—î —Å—É—Ç–Ω—ñ—Å—Ç—å –∫–æ—Ä–∏—Å—Ç—É–≤–∞—á–∞ –≤ —Å–∏—Å—Ç–µ–º—ñ. –í–∏–∫–æ—Ä–∏—Å—Ç–æ–≤—É—î—Ç—å—Å—è MongoDB —è–∫ –±–∞–∑–∞ –¥–∞–Ω–∏—Ö, —Ç–æ–º—É –∫–ª–∞—Å –∞–Ω–æ—Ç—É—î—Ç—å—Å—è @Document. Lombok –∞–Ω–æ—Ç–∞—Ü—ñ—ó (@Data, @Builder) –∞–≤—Ç–æ–º–∞—Ç–∏—á–Ω–æ –≥–µ–Ω–µ—Ä—É—é—Ç—å –≥–µ—Ç—Ç–µ—Ä–∏, —Å–µ—Ç—Ç–µ—Ä–∏ —Ç–∞ builder pattern –¥–ª—è –∑—Ä—É—á–Ω–æ–≥–æ —Å—Ç–≤–æ—Ä–µ–Ω–Ω—è –æ–±\'—î–∫—Ç—ñ–≤.')

add_code_block(doc, [
    'package com.example.cwweb.users;',
    '',
    'import lombok.AllArgsConstructor;',
    'import lombok.Builder;',
    'import lombok.Data;',
    'import lombok.NoArgsConstructor;',
    'import org.springframework.data.annotation.Id;',
    'import org.springframework.data.mongodb.core.mapping.Document;',
    'import java.time.LocalDateTime;',
    '',
    '@Data',
    '@Builder',
    '@NoArgsConstructor',
    '@AllArgsConstructor',
    '@Document(collection = "users")',
    'public class User {',
    '    @Id',
    '    private String id;',
    '    private String email;',
    '    private String passwordHash;',
    '    private String displayName;',
    '    private Role role;',
    '    private Status status;',
    '    private LocalDateTime createdAt;',
    '    private LocalDateTime updatedAt;',
    '',
    '    public enum Role {',
    '        USER, ADMIN',
    '    }',
    '',
    '    public enum Status {',
    '        ACTIVE, BLOCKED',
    '    }',
    '}'
])

add_para(doc, '–ü–æ–ª—è –∫–ª–∞—Å—É: id - —É–Ω—ñ–∫–∞–ª—å–Ω–∏–π —ñ–¥–µ–Ω—Ç–∏—Ñ—ñ–∫–∞—Ç–æ—Ä MongoDB, email - –µ–ª–µ–∫—Ç—Ä–æ–Ω–Ω–∞ –∞–¥—Ä–µ—Å–∞ –∫–æ—Ä–∏—Å—Ç—É–≤–∞—á–∞ (—É–Ω—ñ–∫–∞–ª—å–Ω–∞), passwordHash - —Ö–µ—à –ø–∞—Ä–æ–ª—è (–≤–∏–∫–æ—Ä–∏—Å—Ç–æ–≤—É—î—Ç—å—Å—è BCrypt), displayName - —ñ–º\'—è –¥–ª—è –≤—ñ–¥–æ–±—Ä–∞–∂–µ–Ω–Ω—è, role - —Ä–æ–ª—å —É —Å–∏—Å—Ç–µ–º—ñ (USER –∞–±–æ ADMIN), status - —Å—Ç–∞—Ç—É—Å –∞–∫–∞—É–Ω—Ç—É (ACTIVE –∞–±–æ BLOCKED), createdAt/updatedAt - —á–∞—Å–æ–≤—ñ –º—ñ—Ç–∫–∏ –¥–ª—è –∞—É–¥–∏—Ç—É.')

doc.add_paragraph()

add_explanation(doc,
    '–ö–ª–∞—Å Goal - –º–æ–¥–µ–ª—å —Ü—ñ–ª—ñ –∫–æ—Ä–∏—Å—Ç—É–≤–∞—á–∞:',
    '–ö–ª–∞—Å Goal –ø—Ä–µ–¥—Å—Ç–∞–≤–ª—è—î —Ü—ñ–ª—å –∞–±–æ –∑–≤–∏—á–∫—É, —è–∫—É –∫–æ—Ä–∏—Å—Ç—É–≤–∞—á —Ö–æ—á–µ –≤—ñ–¥—Å—Ç–µ–∂—É–≤–∞—Ç–∏. –ú—ñ—Å—Ç–∏—Ç—å —ñ–Ω—Ñ–æ—Ä–º–∞—Ü—ñ—é –ø—Ä–æ –Ω–∞–∑–≤—É, –æ–ø–∏—Å, —á–∞—Å—Ç–æ—Ç—É –≤–∏–∫–æ–Ω–∞–Ω–Ω—è, —Å—Ç–∞—Ç—É—Å —Ç–∞ —á–∞—Å–æ–≤—ñ —Ä–∞–º–∫–∏. –í–∏–∫–æ—Ä–∏—Å—Ç–æ–≤—É—î—Ç—å—Å—è builder pattern –¥–ª—è –∑—Ä—É—á–Ω–æ–≥–æ —Å—Ç–≤–æ—Ä–µ–Ω–Ω—è —Ü—ñ–ª–µ–π –∑ —Ä—ñ–∑–Ω–∏–º–∏ –ø–∞—Ä–∞–º–µ—Ç—Ä–∞–º–∏.')

add_code_block(doc, [
    'package com.example.cwweb.goals;',
    '',
    'import lombok.AllArgsConstructor;',
    'import lombok.Builder;',
    'import lombok.Data;',
    'import lombok.NoArgsConstructor;',
    'import org.springframework.data.annotation.Id;',
    'import org.springframework.data.mongodb.core.mapping.Document;',
    'import java.time.LocalDate;',
    'import java.time.LocalDateTime;',
    '',
    '@Data',
    '@Builder',
    '@NoArgsConstructor',
    '@AllArgsConstructor',
    '@Document(collection = "goals")',
    'public class Goal {',
    '    @Id',
    '    private String id;',
    '    private String userId;',
    '    private String title;',
    '    private String description;',
    '    private Frequency frequency;',
    '    private LocalDate startDate;',
    '    private LocalDate endDate;',
    '    private boolean isPublic;',
    '    private GoalStatus status;',
    '    private LocalDateTime createdAt;',
    '    private LocalDateTime updatedAt;',
    '}'
])

add_para(doc, '–ü–æ–ª—è –∫–ª–∞—Å—É: userId - –ø–æ—Å–∏–ª–∞–Ω–Ω—è –Ω–∞ –≤–ª–∞—Å–Ω–∏–∫–∞ —Ü—ñ–ª—ñ, title - –Ω–∞–∑–≤–∞ —Ü—ñ–ª—ñ, description - –¥–µ—Ç–∞–ª—å–Ω–∏–π –æ–ø–∏—Å, frequency - —á–∞—Å—Ç–æ—Ç–∞ –≤–∏–∫–æ–Ω–∞–Ω–Ω—è (DAILY, WEEKLY, MONTHLY), startDate/endDate - –ø–µ—Ä—ñ–æ–¥ –¥—ñ—ó —Ü—ñ–ª—ñ, isPublic - –≤–∏–¥–∏–º—ñ—Å—Ç—å –¥–ª—è —ñ–Ω—à–∏—Ö –∫–æ—Ä–∏—Å—Ç—É–≤–∞—á—ñ–≤, status - –ø–æ—Ç–æ—á–Ω–∏–π —Å—Ç–∞—Ç—É—Å (ACTIVE, COMPLETED, PAUSED, ARCHIVED).')

doc.add_paragraph()

add_explanation(doc,
    '–ö–ª–∞—Å Group - –º–æ–¥–µ–ª—å –≥—Ä—É–ø–∏ –¥–ª—è —Å–ø—ñ–ª—å–Ω–æ–≥–æ –≤—ñ–¥—Å—Ç–µ–∂–µ–Ω–Ω—è:',
    '–ö–ª–∞—Å Group –¥–æ–∑–≤–æ–ª—è—î –∫–æ—Ä–∏—Å—Ç—É–≤–∞—á–∞–º –æ–±\'—î–¥–Ω—É–≤–∞—Ç–∏—Å—è –≤ –≥—Ä—É–ø–∏ –¥–ª—è —Å–ø—ñ–ª—å–Ω–æ—ó –º–æ—Ç–∏–≤–∞—Ü—ñ—ó —Ç–∞ –≤—ñ–¥—Å—Ç–µ–∂–µ–Ω–Ω—è –ø—Ä–æ–≥—Ä–µ—Å—É. –ì—Ä—É–ø–∞ –º–∞—î –≤–ª–∞—Å–Ω–∏–∫–∞, –Ω–∞–∑–≤—É, –æ–ø–∏—Å —Ç–∞ –Ω–∞–ª–∞—à—Ç—É–≤–∞–Ω–Ω—è –≤–∏–¥–∏–º–æ—Å—Ç—ñ.')

add_code_block(doc, [
    'package com.example.cwweb.groups;',
    '',
    'import lombok.AllArgsConstructor;',
    'import lombok.Builder;',
    'import lombok.Data;',
    'import lombok.NoArgsConstructor;',
    'import org.springframework.data.annotation.Id;',
    'import org.springframework.data.mongodb.core.mapping.Document;',
    'import java.time.LocalDateTime;',
    '',
    '@Data',
    '@Builder',
    '@NoArgsConstructor',
    '@AllArgsConstructor',
    '@Document(collection = "groups")',
    'public class Group {',
    '    @Id',
    '    private String id;',
    '    private String name;',
    '    private String description;',
    '    private Visibility visibility;',
    '    private String ownerId;',
    '    private LocalDateTime createdAt;',
    '    private LocalDateTime updatedAt;',
    '}'
])

add_para(doc, '–ü–æ–ª—è –∫–ª–∞—Å—É: name - —É–Ω—ñ–∫–∞–ª—å–Ω–∞ –Ω–∞–∑–≤–∞ –≥—Ä—É–ø–∏, visibility - —Ä—ñ–≤–µ–Ω—å –¥–æ—Å—Ç—É–ø—É (PUBLIC –∞–±–æ PRIVATE), ownerId - —ñ–¥–µ–Ω—Ç–∏—Ñ—ñ–∫–∞—Ç–æ—Ä –≤–ª–∞—Å–Ω–∏–∫–∞ –≥—Ä—É–ø–∏, —è–∫–∏–π –º–∞—î —Ä–æ–∑—à–∏—Ä–µ–Ω—ñ –ø—Ä–∞–≤–∞.')

# 4.1.2 Controller Layer
doc.add_page_break()
doc.add_heading('4.1.2 –ö–æ–Ω—Ç—Ä–æ–ª–µ—Ä–∏ - –æ–±—Ä–æ–±–∫–∞ HTTP –∑–∞–ø–∏—Ç—ñ–≤', level=3)

add_explanation(doc,
    '–ö–ª–∞—Å AuthController - –∫–æ–Ω—Ç—Ä–æ–ª–µ—Ä –∞–≤—Ç–µ–Ω—Ç–∏—Ñ—ñ–∫–∞—Ü—ñ—ó:',
    'AuthController –≤—ñ–¥–ø–æ–≤—ñ–¥–∞—î –∑–∞ –æ–±—Ä–æ–±–∫—É –∑–∞–ø–∏—Ç—ñ–≤ —Ä–µ—î—Å—Ç—Ä–∞—Ü—ñ—ó, –∞–≤—Ç–æ—Ä–∏–∑–∞—Ü—ñ—ó —Ç–∞ –æ–Ω–æ–≤–ª–µ–Ω–Ω—è —Ç–æ–∫–µ–Ω—ñ–≤. –í–∏–∫–æ—Ä–∏—Å—Ç–æ–≤—É—î –∞–Ω–æ—Ç–∞—Ü—ñ—é @RestController –¥–ª—è –∞–≤—Ç–æ–º–∞—Ç–∏—á–Ω–æ–≥–æ –ø–µ—Ä–µ—Ç–≤–æ—Ä–µ–Ω–Ω—è –æ–±\'—î–∫—Ç—ñ–≤ —É JSON. –í—Å—ñ endpoint\'–∏ –∑–Ω–∞—Ö–æ–¥—è—Ç—å—Å—è –∑–∞ —à–ª—è—Ö–æ–º /auth.')

add_code_block(doc, [
    'package com.example.cwweb.auth;',
    '',
    'import com.example.cwweb.common.ApiResponse;',
    'import org.springframework.http.ResponseEntity;',
    'import org.springframework.web.bind.annotation.*;',
    '',
    '@RestController',
    '@RequestMapping("/auth")',
    'public class AuthController {',
    '',
    '    private final AuthService authService;',
    '',
    '    public AuthController(AuthService authService) {',
    '        this.authService = authService;',
    '    }',
    '',
    '    @PostMapping("/signup")',
    '    public ResponseEntity<ApiResponse<AuthResponse>> signup(',
    '            @RequestBody SignupRequest request) {',
    '        AuthResponse response = authService.signup(request);',
    '        return ResponseEntity.ok(ApiResponse.<AuthResponse>builder()',
    '                .success(true)',
    '                .message("Signup successful")',
    '                .data(response)',
    '                .build());',
    '    }',
    '',
    '    @PostMapping("/login")',
    '    public ResponseEntity<ApiResponse<AuthResponse>> login(',
    '            @RequestBody LoginRequest request) {',
    '        AuthResponse response = authService.login(request);',
    '        return ResponseEntity.ok(ApiResponse.<AuthResponse>builder()',
    '                .success(true)',
    '                .message("Login successful")',
    '                .data(response)',
    '                .build());',
    '    }',
    '',
    '    @PostMapping("/refresh")',
    '    public ResponseEntity<ApiResponse<AuthResponse>> refresh(',
    '            @RequestBody RefreshRequest request) {',
    '        AuthResponse response = authService.refresh(',
    '                request.getRefreshToken());',
    '        return ResponseEntity.ok(ApiResponse.<AuthResponse>builder()',
    '                .success(true)',
    '                .message("Token refreshed")',
    '                .data(response)',
    '                .build());',
    '    }',
    '}'
])

add_para(doc, '–ú–µ—Ç–æ–¥ signup –æ–±—Ä–æ–±–ª—è—î —Ä–µ—î—Å—Ç—Ä–∞—Ü—ñ—é –Ω–æ–≤–æ–≥–æ –∫–æ—Ä–∏—Å—Ç—É–≤–∞—á–∞, –ø—Ä–∏–π–º–∞—é—á–∏ email, –ø–∞—Ä–æ–ª—å —Ç–∞ displayName. –ú–µ—Ç–æ–¥ login –≤–∏–∫–æ–Ω—É—î –∞–≤—Ç–æ—Ä–∏–∑–∞—Ü—ñ—é –∫–æ—Ä–∏—Å—Ç—É–≤–∞—á–∞ —Ç–∞ –ø–æ–≤–µ—Ä—Ç–∞—î JWT —Ç–æ–∫–µ–Ω–∏. –ú–µ—Ç–æ–¥ refresh –æ–Ω–æ–≤–ª—é—î access token –∑–∞ –¥–æ–ø–æ–º–æ–≥–æ—é refresh token. –í—Å—ñ –≤—ñ–¥–ø–æ–≤—ñ–¥—ñ –∑–∞–≥–æ—Ä—Ç–∞—é—Ç—å—Å—è –≤ —É–Ω—ñ—Ñ—ñ–∫–æ–≤–∞–Ω–∏–π ApiResponse –¥–ª—è –∫–æ–Ω—Å–∏—Å—Ç–µ–Ω—Ç–Ω–æ—Å—Ç—ñ API.')

doc.add_paragraph()

add_explanation(doc,
    '–ö–ª–∞—Å GoalController - –∫–æ–Ω—Ç—Ä–æ–ª–µ—Ä —É–ø—Ä–∞–≤–ª—ñ–Ω–Ω—è —Ü—ñ–ª—è–º–∏:',
    'GoalController –Ω–∞–¥–∞—î REST API –¥–ª—è CRUD –æ–ø–µ—Ä–∞—Ü—ñ–π –Ω–∞–¥ —Ü—ñ–ª—è–º–∏. –í–∏–∫–æ—Ä–∏—Å—Ç–æ–≤—É—î @AuthenticationPrincipal –¥–ª—è –æ—Ç—Ä–∏–º–∞–Ω–Ω—è —ñ–Ω—Ñ–æ—Ä–º–∞—Ü—ñ—ó –ø—Ä–æ –∞–≤—Ç–æ—Ä–∏–∑–æ–≤–∞–Ω–æ–≥–æ –∫–æ—Ä–∏—Å—Ç—É–≤–∞—á–∞ –∑ JWT —Ç–æ–∫–µ–Ω–∞. –í—Å—ñ –º–µ—Ç–æ–¥–∏ –∑–∞—Ö–∏—â–µ–Ω—ñ –∞–≤—Ç–µ–Ω—Ç–∏—Ñ—ñ–∫–∞—Ü—ñ—î—é.')

add_code_block(doc, [
    'package com.example.cwweb.goals;',
    '',
    'import com.example.cwweb.common.ApiResponse;',
    'import org.springframework.http.ResponseEntity;',
    'import org.springframework.security.core.annotation.AuthenticationPrincipal;',
    'import org.springframework.security.core.userdetails.UserDetails;',
    'import org.springframework.web.bind.annotation.*;',
    'import java.util.List;',
    '',
    '@RestController',
    '@RequestMapping("/goals")',
    'public class GoalController {',
    '    private final GoalService goalService;',
    '',
    '    public GoalController(GoalService goalService) {',
    '        this.goalService = goalService;',
    '    }',
    '',
    '    @PostMapping',
    '    public ResponseEntity<ApiResponse<Goal>> createGoal(',
    '            @AuthenticationPrincipal UserDetails userDetails,',
    '            @RequestBody CreateGoalRequest request) {',
    '        String userId = userDetails.getUsername();',
    '        Goal goal = goalService.createGoal(userId, request);',
    '        return ResponseEntity.ok(ApiResponse.<Goal>builder()',
    '                .success(true)',
    '                .message("Goal created")',
    '                .data(goal)',
    '                .build());',
    '    }',
    '',
    '    @GetMapping',
    '    public ResponseEntity<ApiResponse<List<Goal>>> getUserGoals(',
    '            @AuthenticationPrincipal UserDetails userDetails) {',
    '        String userId = userDetails.getUsername();',
    '        List<Goal> goals = goalService.getUserGoals(userId);',
    '        return ResponseEntity.ok(ApiResponse.<List<Goal>>builder()',
    '                .success(true)',
    '                .data(goals)',
    '                .build());',
    '    }',
    '',
    '    @PutMapping("/{id}")',
    '    public ResponseEntity<ApiResponse<Goal>> updateGoal(',
    '            @PathVariable String id,',
    '            @AuthenticationPrincipal UserDetails userDetails,',
    '            @RequestBody UpdateGoalRequest request) {',
    '        String userId = userDetails.getUsername();',
    '        Goal goal = goalService.updateGoal(id, userId, request);',
    '        return ResponseEntity.ok(ApiResponse.<Goal>builder()',
    '                .success(true)',
    '                .message("Goal updated")',
    '                .data(goal)',
    '                .build());',
    '    }',
    '',
    '    @DeleteMapping("/{id}")',
    '    public ResponseEntity<ApiResponse<Void>> deleteGoal(',
    '            @PathVariable String id,',
    '            @AuthenticationPrincipal UserDetails userDetails) {',
    '        String userId = userDetails.getUsername();',
    '        goalService.deleteGoal(id, userId);',
    '        return ResponseEntity.ok(ApiResponse.<Void>builder()',
    '                .success(true)',
    '                .message("Goal deleted")',
    '                .build());',
    '    }',
    '}'
])

add_para(doc, '–ö–æ–Ω—Ç—Ä–æ–ª–µ—Ä –Ω–∞–¥–∞—î 5 –æ—Å–Ω–æ–≤–Ω–∏—Ö endpoint\'—ñ–≤: POST /goals - —Å—Ç–≤–æ—Ä–µ–Ω–Ω—è –Ω–æ–≤–æ—ó —Ü—ñ–ª—ñ, GET /goals - –æ—Ç—Ä–∏–º–∞–Ω–Ω—è –≤—Å—ñ—Ö —Ü—ñ–ª–µ–π –∫–æ—Ä–∏—Å—Ç—É–≤–∞—á–∞, GET /goals/{id} - –æ—Ç—Ä–∏–º–∞–Ω–Ω—è –∫–æ–Ω–∫—Ä–µ—Ç–Ω–æ—ó —Ü—ñ–ª—ñ, PUT /goals/{id} - –æ–Ω–æ–≤–ª–µ–Ω–Ω—è —Ü—ñ–ª—ñ, DELETE /goals/{id} - –≤–∏–¥–∞–ª–µ–Ω–Ω—è —Ü—ñ–ª—ñ. –ê–≤—Ç–æ–º–∞—Ç–∏—á–Ω–∞ –≤–∞–ª—ñ–¥–∞—Ü—ñ—è –¥–æ—Å—Ç—É–ø—É –∑–∞–±–µ–∑–ø–µ—á—É—î—Ç—å—Å—è –ø–µ—Ä–µ–¥–∞—á–µ—é userId –¥–æ —Å–µ—Ä–≤—ñ—Å–Ω–æ–≥–æ —à–∞—Ä—É.')

doc.add_paragraph()

add_explanation(doc,
    '–ö–ª–∞—Å ProgressController - –∫–æ–Ω—Ç—Ä–æ–ª–µ—Ä –ª–æ–≥—É–≤–∞–Ω–Ω—è –ø—Ä–æ–≥—Ä–µ—Å—É:',
    'ProgressController –¥–æ–∑–≤–æ–ª—è—î –∫–æ—Ä–∏—Å—Ç—É–≤–∞—á–∞–º –ª–æ–≥—É–≤–∞—Ç–∏ —â–æ–¥–µ–Ω–Ω–∏–π –ø—Ä–æ–≥—Ä–µ—Å –ø–æ —Ü—ñ–ª—è—Ö —Ç–∞ –æ—Ç—Ä–∏–º—É–≤–∞—Ç–∏ —ñ—Å—Ç–æ—Ä—ñ—é –≤–∏–∫–æ–Ω–∞–Ω–Ω—è. –ö–æ–∂–µ–Ω –∑–∞–ø–∏—Å –ø—Ä–æ–≥—Ä–µ—Å—É –∑–±–µ—Ä—ñ–≥–∞—î—Ç—å—Å—è –∑ –¥–∞—Ç–æ—é —Ç–∞ –æ–ø—Ü—ñ–æ–Ω–∞–ª—å–Ω–∏–º–∏ –Ω–æ—Ç–∞—Ç–∫–∞–º–∏.')

add_code_block(doc, [
    'package com.example.cwweb.progress;',
    '',
    'import com.example.cwweb.common.ApiResponse;',
    'import org.springframework.http.ResponseEntity;',
    'import org.springframework.security.core.annotation.AuthenticationPrincipal;',
    'import org.springframework.security.core.userdetails.UserDetails;',
    'import org.springframework.web.bind.annotation.*;',
    'import java.util.List;',
    '',
    '@RestController',
    '@RequestMapping("/progress")',
    'public class ProgressController {',
    '    private final ProgressService progressService;',
    '',
    '    public ProgressController(ProgressService progressService) {',
    '        this.progressService = progressService;',
    '    }',
    '',
    '    @PostMapping',
    '    public ResponseEntity<ApiResponse<ProgressLog>> logProgress(',
    '            @AuthenticationPrincipal UserDetails userDetails,',
    '            @RequestBody LogProgressRequest request) {',
    '        String userId = userDetails.getUsername();',
    '        ProgressLog log = progressService.logProgress(userId, request);',
    '        return ResponseEntity.ok(ApiResponse.<ProgressLog>builder()',
    '                .success(true)',
    '                .message("Progress logged")',
    '                .data(log)',
    '                .build());',
    '    }',
    '',
    '    @GetMapping("/goal/{goalId}")',
    '    public ResponseEntity<ApiResponse<List<ProgressLog>>> getGoalProgress(',
    '            @PathVariable String goalId,',
    '            @AuthenticationPrincipal UserDetails userDetails) {',
    '        String userId = userDetails.getUsername();',
    '        List<ProgressLog> logs = progressService.getGoalProgress(',
    '                goalId, userId);',
    '        return ResponseEntity.ok(ApiResponse.<List<ProgressLog>>builder()',
    '                .success(true)',
    '                .data(logs)',
    '                .build());',
    '    }',
    '}'
])

add_para(doc, 'Endpoint POST /progress –ø—Ä–∏–π–º–∞—î goalId, –¥–∞—Ç—É —Ç–∞ –æ–ø—Ü—ñ–æ–Ω–∞–ª—å–Ω—ñ –Ω–æ—Ç–∞—Ç–∫–∏ –¥–ª—è —Å—Ç–≤–æ—Ä–µ–Ω–Ω—è –∑–∞–ø–∏—Å—É –ø—Ä–æ–≥—Ä–µ—Å—É. GET /progress/goal/{goalId} –ø–æ–≤–µ—Ä—Ç–∞—î –≤—Å—é —ñ—Å—Ç–æ—Ä—ñ—é –≤–∏–∫–æ–Ω–∞–Ω–Ω—è –∫–æ–Ω–∫—Ä–µ—Ç–Ω–æ—ó —Ü—ñ–ª—ñ —É —Ö—Ä–æ–Ω–æ–ª–æ–≥—ñ—á–Ω–æ–º—É –ø–æ—Ä—è–¥–∫—É.')

# 4.1.3 Service Layer
doc.add_page_break()
doc.add_heading('4.1.3 –°–µ—Ä–≤—ñ—Å–Ω–∏–π —à–∞—Ä - –±—ñ–∑–Ω–µ—Å-–ª–æ–≥—ñ–∫–∞', level=3)

add_explanation(doc,
    '–ö–ª–∞—Å GoalService - —Å–µ—Ä–≤—ñ—Å —É–ø—Ä–∞–≤–ª—ñ–Ω–Ω—è —Ü—ñ–ª—è–º–∏:',
    'GoalService –º—ñ—Å—Ç–∏—Ç—å –±—ñ–∑–Ω–µ—Å-–ª–æ–≥—ñ–∫—É —Ä–æ–±–æ—Ç–∏ –∑ —Ü—ñ–ª—è–º–∏. –í–∏–∫–æ–Ω—É—î –≤–∞–ª—ñ–¥–∞—Ü—ñ—é –¥–∞–Ω–∏—Ö, –ø–µ—Ä–µ–≤—ñ—Ä–∫—É –ø—Ä–∞–≤ –¥–æ—Å—Ç—É–ø—É —Ç–∞ –≤–∑–∞—î–º–æ–¥—ñ—é –∑ —Ä–µ–ø–æ–∑–∏—Ç–æ—Ä—ñ—î–º. –í–∏–∫–æ—Ä–∏—Å—Ç–æ–≤—É—î pattern Builder –¥–ª—è —Å—Ç–≤–æ—Ä–µ–Ω–Ω—è –æ–±\'—î–∫—Ç—ñ–≤ Goal.')

add_code_block(doc, [
    'package com.example.cwweb.goals;',
    '',
    'import com.example.cwweb.common.NotFoundException;',
    'import com.example.cwweb.common.ForbiddenException;',
    'import org.springframework.stereotype.Service;',
    'import java.time.LocalDateTime;',
    'import java.util.List;',
    '',
    '@Service',
    'public class GoalService {',
    '    private final GoalRepository goalRepository;',
    '',
    '    public GoalService(GoalRepository goalRepository) {',
    '        this.goalRepository = goalRepository;',
    '    }',
    '',
    '    public Goal createGoal(String userId, CreateGoalRequest request) {',
    '        Goal goal = Goal.builder()',
    '                .userId(userId)',
    '                .title(request.getTitle())',
    '                .description(request.getDescription())',
    '                .frequency(request.getFrequency())',
    '                .startDate(request.getStartDate())',
    '                .endDate(request.getEndDate())',
    '                .isPublic(request.isPublic())',
    '                .status(GoalStatus.ACTIVE)',
    '                .createdAt(LocalDateTime.now())',
    '                .updatedAt(LocalDateTime.now())',
    '                .build();',
    '        return goalRepository.save(goal);',
    '    }',
    '',
    '    public List<Goal> getUserGoals(String userId) {',
    '        return goalRepository.findByUserId(userId);',
    '    }',
    '',
    '    public Goal getGoal(String goalId, String userId) {',
    '        Goal goal = goalRepository.findById(goalId)',
    '                .orElseThrow(() -> new NotFoundException("Goal not found"));',
    '        if (!goal.getUserId().equals(userId) && !goal.isPublic()) {',
    '            throw new ForbiddenException("Access denied");',
    '        }',
    '        return goal;',
    '    }',
    '',
    '    public Goal updateGoal(String goalId, String userId,',
    '                          UpdateGoalRequest request) {',
    '        Goal goal = goalRepository.findById(goalId)',
    '                .orElseThrow(() -> new NotFoundException("Goal not found"));',
    '        if (!goal.getUserId().equals(userId)) {',
    '            throw new ForbiddenException("Access denied");',
    '        }',
    '        if (request.getTitle() != null)',
    '            goal.setTitle(request.getTitle());',
    '        if (request.getDescription() != null)',
    '            goal.setDescription(request.getDescription());',
    '        if (request.getStatus() != null)',
    '            goal.setStatus(request.getStatus());',
    '        goal.setUpdatedAt(LocalDateTime.now());',
    '        return goalRepository.save(goal);',
    '    }',
    '',
    '    public void deleteGoal(String goalId, String userId) {',
    '        Goal goal = goalRepository.findById(goalId)',
    '                .orElseThrow(() -> new NotFoundException("Goal not found"));',
    '        if (!goal.getUserId().equals(userId)) {',
    '            throw new ForbiddenException("Access denied");',
    '        }',
    '        goalRepository.deleteById(goalId);',
    '    }',
    '}'
])

add_para(doc, '–°–µ—Ä–≤—ñ—Å –∑–∞–±–µ–∑–ø–µ—á—É—î –±–µ–∑–ø–µ–∫—É —á–µ—Ä–µ–∑ –ø–µ—Ä–µ–≤—ñ—Ä–∫—É userId –ø—Ä–∏ –∫–æ–∂–Ω—ñ–π –æ–ø–µ—Ä–∞—Ü—ñ—ó. –ú–µ—Ç–æ–¥ createGoal –∞–≤—Ç–æ–º–∞—Ç–∏—á–Ω–æ –≤—Å—Ç–∞–Ω–æ–≤–ª—é—î —Å—Ç–∞—Ç—É—Å ACTIVE —Ç–∞ –ø–æ—Ç–æ—á–Ω–∏–π —á–∞—Å —Å—Ç–≤–æ—Ä–µ–Ω–Ω—è. –ú–µ—Ç–æ–¥ getGoal –¥–æ–∑–≤–æ–ª—è—î –¥–æ—Å—Ç—É–ø –¥–æ –ø—É–±–ª—ñ—á–Ω–∏—Ö —Ü—ñ–ª–µ–π —ñ–Ω—à–∏—Ö –∫–æ—Ä–∏—Å—Ç—É–≤–∞—á—ñ–≤. –ú–µ—Ç–æ–¥ updateGoal –≤–∏–∫–æ—Ä–∏—Å—Ç–æ–≤—É—î –ø–∞—Ç–µ—Ä–Ω —á–∞—Å—Ç–∫–æ–≤–æ–≥–æ –æ–Ω–æ–≤–ª–µ–Ω–Ω—è (patch), –¥–æ–∑–≤–æ–ª—è—é—á–∏ –∑–º—ñ–Ω—é–≤–∞—Ç–∏ –ª–∏—à–µ –ø–µ—Ä–µ–¥–∞–Ω—ñ –ø–æ–ª—è.')

# 4.2 Frontend —Ä–µ–∞–ª—ñ–∑–∞—Ü—ñ—è
doc.add_page_break()
doc.add_heading('4.2 –†–µ–∞–ª—ñ–∑–∞—Ü—ñ—è frontend —á–∞—Å—Ç–∏–Ω–∏ —Å–∏—Å—Ç–µ–º–∏', level=2)

doc.add_heading('4.2.1 –°—Ç—Ä—É–∫—Ç—É—Ä–∞ React –¥–æ–¥–∞—Ç–∫—É', level=3)

add_explanation(doc,
    '–§–∞–π–ª App.jsx - –≥–æ–ª–æ–≤–Ω–∏–π –∫–æ–º–ø–æ–Ω–µ–Ω—Ç –¥–æ–¥–∞—Ç–∫—É:',
    'App.jsx —î —Ç–æ—á–∫–æ—é –≤—Ö–æ–¥—É –¥–æ–¥–∞—Ç–∫—É. –ù–∞–ª–∞—à—Ç–æ–≤—É—î –º–∞—Ä—à—Ä—É—Ç–∏–∑–∞—Ü—ñ—é –∑–∞ –¥–æ–ø–æ–º–æ–≥–æ—é React Router v6, –≤–∏–∑–Ω–∞—á–∞—î –∑–∞—Ö–∏—â–µ–Ω—ñ —Ç–∞ –ø—É–±–ª—ñ—á–Ω—ñ –º–∞—Ä—à—Ä—É—Ç–∏, –ø—ñ–¥–∫–ª—é—á–∞—î —Å–∏—Å—Ç–µ–º—É —Å–ø–æ–≤—ñ—â–µ–Ω—å (react-hot-toast). –í–∏–∫–æ—Ä–∏—Å—Ç–æ–≤—É—î Zustand store –¥–ª—è —É–ø—Ä–∞–≤–ª—ñ–Ω–Ω—è —Å—Ç–∞–Ω–æ–º –∞–≤—Ç–µ–Ω—Ç–∏—Ñ—ñ–∫–∞—Ü—ñ—ó.')

add_code_block(doc, [
    "import { BrowserRouter, Routes, Route, Navigate }",
    "        from 'react-router-dom'",
    "import { Toaster } from 'react-hot-toast'",
    "import { useAuthStore } from './store/authStore'",
    "import Layout from './components/Layout'",
    "import Login from './pages/Login'",
    "import Dashboard from './pages/Dashboard'",
    "import Goals from './pages/Goals'",
    "",
    "function ProtectedRoute({ children }) {",
    "  const { token } = useAuthStore()",
    "  return token ? children : <Navigate to='/login' />",
    "}",
    "",
    "function PublicRoute({ children }) {",
    "  const { token } = useAuthStore()",
    "  return !token ? children : <Navigate to='/dashboard' />",
    "}",
    "",
    "function App() {",
    "  return (",
    "    <BrowserRouter>",
    "      <Toaster position='top-right' />",
    "      <Routes>",
    "        <Route path='/login' element={",
    "          <PublicRoute><Login /></PublicRoute>",
    "        } />",
    "        <Route path='/' element={",
    "          <ProtectedRoute><Layout /></ProtectedRoute>",
    "        }>",
    "          <Route index element={<Navigate to='/dashboard' />} />",
    "          <Route path='dashboard' element={<Dashboard />} />",
    "          <Route path='goals' element={<Goals />} />",
    "        </Route>",
    "      </Routes>",
    "    </BrowserRouter>",
    "  )",
    "}",
    "",
    "export default App"
], "JavaScript")

add_para(doc, '–ö–æ–º–ø–æ–Ω–µ–Ω—Ç–∏ ProtectedRoute —Ç–∞ PublicRoute –∑–∞–±–µ–∑–ø–µ—á—É—é—Ç—å –∫–æ–Ω—Ç—Ä–æ–ª—å –¥–æ—Å—Ç—É–ø—É: ProtectedRoute –ø–µ—Ä–µ–Ω–∞–ø—Ä–∞–≤–ª—è—î –Ω–µ–∞–≤—Ç–æ—Ä–∏–∑–æ–≤–∞–Ω–∏—Ö –∫–æ—Ä–∏—Å—Ç—É–≤–∞—á—ñ–≤ –Ω–∞ /login, PublicRoute –ø–µ—Ä–µ–Ω–∞–ø—Ä–∞–≤–ª—è—î –∞–≤—Ç–æ—Ä–∏–∑–æ–≤–∞–Ω–∏—Ö –Ω–∞ /dashboard. –í–∏–∫–æ—Ä–∏—Å—Ç–∞–Ω–Ω—è –≤–∫–ª–∞–¥–µ–Ω–∏—Ö –º–∞—Ä—à—Ä—É—Ç—ñ–≤ (nested routes) –¥–æ–∑–≤–æ–ª—è—î Layout –±—É—Ç–∏ –æ–±–≥–æ—Ä—Ç–∫–æ—é –¥–ª—è –≤—Å—ñ—Ö –∑–∞—Ö–∏—â–µ–Ω–∏—Ö —Å—Ç–æ—Ä—ñ–Ω–æ–∫.')

doc.add_paragraph()

doc.add_heading('4.2.2 State Management - Zustand', level=3)

add_explanation(doc,
    '–§–∞–π–ª authStore.js - –≥–ª–æ–±–∞–ª—å–Ω–∏–π —Å—Ç–∞–Ω –∞–≤—Ç–µ–Ω—Ç–∏—Ñ—ñ–∫–∞—Ü—ñ—ó:',
    'Zustand store —É–ø—Ä–∞–≤–ª—è—î —Å—Ç–∞–Ω–æ–º –∞–≤—Ç–µ–Ω—Ç–∏—Ñ—ñ–∫–∞—Ü—ñ—ó –∫–æ—Ä–∏—Å—Ç—É–≤–∞—á–∞. –í–∏–∫–æ—Ä–∏—Å—Ç–æ–≤—É—î middleware persist –¥–ª—è –∑–±–µ—Ä–µ–∂–µ–Ω–Ω—è —Ç–æ–∫–µ–Ω—ñ–≤ —É localStorage, —â–æ –¥–æ–∑–≤–æ–ª—è—î –∑–±–µ—Ä—ñ–≥–∞—Ç–∏ —Å–µ—Å—ñ—é –ø—ñ—Å–ª—è –ø–µ—Ä–µ–∑–∞–≤–∞–Ω—Ç–∞–∂–µ–Ω–Ω—è —Å—Ç–æ—Ä—ñ–Ω–∫–∏. –ù–∞–¥–∞—î –º–µ—Ç–æ–¥–∏ –¥–ª—è login, logout —Ç–∞ –æ–Ω–æ–≤–ª–µ–Ω–Ω—è —Ç–æ–∫–µ–Ω—ñ–≤.')

add_code_block(doc, [
    "import { create } from 'zustand'",
    "import { persist } from 'zustand/middleware'",
    "",
    "export const useAuthStore = create(",
    "  persist(",
    "    (set) => ({",
    "      token: null,",
    "      refreshToken: null,",
    "      user: null,",
    "      ",
    "      setTokens: (accessToken, refreshToken) => {",
    "        set({ token: accessToken, refreshToken })",
    "      },",
    "      ",
    "      setUser: (user) => {",
    "        set({ user })",
    "      },",
    "      ",
    "      login: (accessToken, refreshToken, user) => {",
    "        set({ token: accessToken, refreshToken, user })",
    "      },",
    "      ",
    "      logout: () => {",
    "        set({ token: null, refreshToken: null, user: null })",
    "      },",
    "    }),",
    "    {",
    "      name: 'auth-storage',",
    "    }",
    "  )",
    ")"
], "JavaScript")

add_para(doc, 'Store –∑–±–µ—Ä—ñ–≥–∞—î —Ç—Ä–∏ –æ—Å–Ω–æ–≤–Ω—ñ –ø–æ–ª—è: token (JWT access token), refreshToken (–¥–ª—è –æ–Ω–æ–≤–ª–µ–Ω–Ω—è —Ç–æ–∫–µ–Ω—ñ–≤), user (–¥–∞–Ω—ñ –∫–æ—Ä–∏—Å—Ç—É–≤–∞—á–∞). –ú–µ—Ç–æ–¥ login –≤—Å—Ç–∞–Ω–æ–≤–ª—é—î –≤—Å—ñ —Ç—Ä–∏ –∑–Ω–∞—á–µ–Ω–Ω—è –æ–¥–Ω–æ—á–∞—Å–Ω–æ –ø—Ä–∏ —É—Å–ø—ñ—à–Ω—ñ–π –∞–≤—Ç–æ—Ä–∏–∑–∞—Ü—ñ—ó. –ú–µ—Ç–æ–¥ logout –æ—á–∏—â–∞—î —Å—Ç–∞–Ω. Persist middleware –∞–≤—Ç–æ–º–∞—Ç–∏—á–Ω–æ —Å–∏–Ω—Ö—Ä–æ–Ω—ñ–∑—É—î –∑–º—ñ–Ω–∏ –∑ localStorage.')

doc.add_paragraph()

doc.add_heading('4.2.3 API –∫–ª—ñ—î–Ω—Ç–∏', level=3)

add_explanation(doc,
    '–§–∞–π–ª client.js - –Ω–∞–ª–∞—à—Ç—É–≤–∞–Ω–Ω—è Axios:',
    '–ë–∞–∑–æ–≤–∏–π HTTP –∫–ª—ñ—î–Ω—Ç –Ω–∞ –æ—Å–Ω–æ–≤—ñ Axios –∑ –Ω–∞–ª–∞—à—Ç–æ–≤–∞–Ω–∏–º–∏ interceptor\'–∞–º–∏. Request interceptor –¥–æ–¥–∞—î JWT —Ç–æ–∫–µ–Ω –¥–æ –∫–æ–∂–Ω–æ–≥–æ –∑–∞–ø–∏—Ç—É. Response interceptor –æ–±—Ä–æ–±–ª—è—î –ø–æ–º–∏–ª–∫–∏ 401 —Ç–∞ –∞–≤—Ç–æ–º–∞—Ç–∏—á–Ω–æ –æ–Ω–æ–≤–ª—é—î —Ç–æ–∫–µ–Ω–∏ —á–µ—Ä–µ–∑ refresh endpoint.')

add_code_block(doc, [
    "import axios from 'axios'",
    "import { useAuthStore } from '../store/authStore'",
    "",
    "const client = axios.create({",
    "  baseURL: '/api',",
    "  headers: { 'Content-Type': 'application/json' },",
    "})",
    "",
    "client.interceptors.request.use((config) => {",
    "  const token = useAuthStore.getState().token",
    "  if (token) {",
    "    config.headers.Authorization = `Bearer ${token}`",
    "  }",
    "  return config",
    "})",
    "",
    "client.interceptors.response.use(",
    "  (response) => response,",
    "  async (error) => {",
    "    const originalRequest = error.config",
    "    if (error.response?.status === 401 && !originalRequest._retry) {",
    "      originalRequest._retry = true",
    "      try {",
    "        const refreshToken = useAuthStore.getState().refreshToken",
    "        const { data } = await axios.post('/api/auth/refresh',",
    "                                          { refreshToken })",
    "        useAuthStore.getState().setTokens(",
    "          data.data.accessToken, data.data.refreshToken",
    "        )",
    "        originalRequest.headers.Authorization =",
    "          `Bearer ${data.data.accessToken}`",
    "        return client(originalRequest)",
    "      } catch (refreshError) {",
    "        useAuthStore.getState().logout()",
    "        window.location.href = '/login'",
    "        return Promise.reject(refreshError)",
    "      }",
    "    }",
    "    return Promise.reject(error)",
    "  }",
    ")",
    "",
    "export default client"
], "JavaScript")

add_para(doc, 'Response interceptor —Ä–µ–∞–ª—ñ–∑—É—î –∞–≤—Ç–æ–º–∞—Ç–∏—á–Ω–µ –æ–Ω–æ–≤–ª–µ–Ω–Ω—è —Ç–æ–∫–µ–Ω—ñ–≤: –ø—Ä–∏ –æ—Ç—Ä–∏–º–∞–Ω–Ω—ñ 401 –ø–æ–º–∏–ª–∫–∏ —Ä–æ–±–∏—Ç—å—Å—è —Å–ø—Ä–æ–±–∞ –æ–Ω–æ–≤–∏—Ç–∏ —Ç–æ–∫–µ–Ω —á–µ—Ä–µ–∑ /auth/refresh, —è–∫—â–æ —É—Å–ø—ñ—à–Ω–æ - –ø–æ–≤—Ç–æ—Ä—é—î—Ç—å—Å—è –æ—Ä–∏–≥—ñ–Ω–∞–ª—å–Ω–∏–π –∑–∞–ø–∏—Ç –∑ –Ω–æ–≤–∏–º —Ç–æ–∫–µ–Ω–æ–º, —è–∫—â–æ –Ω—ñ - –∫–æ—Ä–∏—Å—Ç—É–≤–∞—á —Ä–æ–∑–ª–æ–≥—ñ–Ω—é—î—Ç—å—Å—è. –ü—Ä–∞–ø–æ—Ä–µ—Ü—å _retry –∑–∞–ø–æ–±—ñ–≥–∞—î –Ω–µ—Å–∫—ñ–Ω—á–µ–Ω–Ω–∏–º —Ü–∏–∫–ª–∞–º.')

doc.add_paragraph()

add_explanation(doc,
    '–§–∞–π–ª goals.js - API –º–µ—Ç–æ–¥–∏ –¥–ª—è —Ä–æ–±–æ—Ç–∏ –∑ —Ü—ñ–ª—è–º–∏:',
    '–ú–æ–¥—É–ª—å goalsAPI —ñ–Ω–∫–∞–ø—Å—É–ª—é—î –≤—Å—ñ HTTP –∑–∞–ø–∏—Ç–∏ –¥–ª—è —Ä–æ–±–æ—Ç–∏ –∑ —Ü—ñ–ª—è–º–∏. –ö–æ–∂–µ–Ω –º–µ—Ç–æ–¥ –ø–æ–≤–µ—Ä—Ç–∞—î Promise –∑ —Ä–æ–∑–ø–∞–∫–æ–≤–∞–Ω–∏–º–∏ –¥–∞–Ω–∏–º–∏ –∑ ApiResponse. –í–∏–∫–æ—Ä–∏—Å—Ç–æ–≤—É—î –±–∞–∑–æ–≤–∏–π client –∑ –∞–≤—Ç–æ–º–∞—Ç–∏—á–Ω–æ—é –∞–≤—Ç–µ–Ω—Ç–∏—Ñ—ñ–∫–∞—Ü—ñ—î—é.')

add_code_block(doc, [
    "import client from './client'",
    "",
    "export const goalsAPI = {",
    "  getAll: async () => {",
    "    const { data } = await client.get('/goals')",
    "    return data.data",
    "  },",
    "  ",
    "  getById: async (id) => {",
    "    const { data } = await client.get(`/goals/${id}`)",
    "    return data.data",
    "  },",
    "  ",
    "  create: async (goalData) => {",
    "    const { data } = await client.post('/goals', goalData)",
    "    return data.data",
    "  },",
    "  ",
    "  update: async (id, goalData) => {",
    "    const { data } = await client.put(`/goals/${id}`, goalData)",
    "    return data.data",
    "  },",
    "  ",
    "  delete: async (id) => {",
    "    await client.delete(`/goals/${id}`)",
    "  },",
    "  ",
    "  logProgress: async (progressData) => {",
    "    const { data } = await client.post('/progress', progressData)",
    "    return data.data",
    "  },",
    "}"
], "JavaScript")

add_para(doc, 'API –º–µ—Ç–æ–¥–∏: getAll –æ—Ç—Ä–∏–º—É—î —Å–ø–∏—Å–æ–∫ –≤—Å—ñ—Ö —Ü—ñ–ª–µ–π –∫–æ—Ä–∏—Å—Ç—É–≤–∞—á–∞, getById –æ—Ç—Ä–∏–º—É—î –∫–æ–Ω–∫—Ä–µ—Ç–Ω—É —Ü—ñ–ª—å, create —Å—Ç–≤–æ—Ä—é—î –Ω–æ–≤—É —Ü—ñ–ª—å, update –æ–Ω–æ–≤–ª—é—î —ñ—Å–Ω—É—é—á—É (patch), delete –≤–∏–¥–∞–ª—è—î —Ü—ñ–ª—å, logProgress –ª–æ–≥—É—î –ø—Ä–æ–≥—Ä–µ—Å –≤–∏–∫–æ–Ω–∞–Ω–Ω—è. –í—Å—ñ –º–µ—Ç–æ–¥–∏ –∞–≤—Ç–æ–º–∞—Ç–∏—á–Ω–æ –¥–æ–¥–∞—é—Ç—å JWT —Ç–æ–∫–µ–Ω —á–µ—Ä–µ–∑ interceptor.')

doc.add_page_break()
doc.add_heading('4.2.4 React –∫–æ–º–ø–æ–Ω–µ–Ω—Ç–∏ —Ç–∞ —Å—Ç–æ—Ä—ñ–Ω–∫–∏', level=3)

add_explanation(doc,
    '–ö–æ–º–ø–æ–Ω–µ–Ω—Ç Layout.jsx - –∑–∞–≥–∞–ª—å–Ω–∞ —Å—Ç—Ä—É–∫—Ç—É—Ä–∞ –¥–æ–¥–∞—Ç–∫—É:',
    'Layout –∫–æ–º–ø–æ–Ω–µ–Ω—Ç –∑–∞–±–µ–∑–ø–µ—á—É—î –∫–æ–Ω—Å–∏—Å—Ç–µ–Ω—Ç–Ω—É —Å—Ç—Ä—É–∫—Ç—É—Ä—É –≤—Å—ñ—Ö —Å—Ç–æ—Ä—ñ–Ω–æ–∫. –ú—ñ—Å—Ç–∏—Ç—å –±—ñ—á–Ω—É –ø–∞–Ω–µ–ª—å –Ω–∞–≤—ñ–≥–∞—Ü—ñ—ó –∑ —ñ–∫–æ–Ω–∫–∞–º–∏ (–≤–∏–∫–æ—Ä–∏—Å—Ç–æ–≤—É—é—á–∏ lucide-react), —à–∞–ø–∫—É –∑ —ñ–Ω—Ñ–æ—Ä–º–∞—Ü—ñ—î—é –∫–æ—Ä–∏—Å—Ç—É–≤–∞—á–∞ —Ç–∞ –æ—Å–Ω–æ–≤–Ω—É –æ–±–ª–∞—Å—Ç—å –∫–æ–Ω—Ç–µ–Ω—Ç—É. –í–∏–∫–æ—Ä–∏—Å—Ç–æ–≤—É—î Outlet –≤—ñ–¥ React Router –¥–ª—è —Ä–µ–Ω–¥–µ—Ä–∏–Ω–≥—É –≤–∫–ª–∞–¥–µ–Ω–∏—Ö –º–∞—Ä—à—Ä—É—Ç—ñ–≤.')

add_code_block(doc, [
    "import { Outlet, NavLink, useNavigate } from 'react-router-dom'",
    "import { useAuthStore } from '../store/authStore'",
    "import { Home, Target, Users, Award, LogOut } from 'lucide-react'",
    "",
    "export default function Layout() {",
    "  const { user, logout } = useAuthStore()",
    "  const navigate = useNavigate()",
    "",
    "  const handleLogout = () => {",
    "    logout()",
    "    navigate('/login')",
    "  }",
    "",
    "  const navItems = [",
    "    { to: '/dashboard', icon: Home, label: 'Dashboard' },",
    "    { to: '/goals', icon: Target, label: 'Goals' },",
    "    { to: '/groups', icon: Users, label: 'Groups' },",
    "    { to: '/achievements', icon: Award, label: 'Achievements' },",
    "  ]",
    "",
    "  return (",
    "    <div className='flex h-screen bg-gray-100'>",
    "      <aside className='w-64 bg-white shadow-lg'>",
    "        <div className='p-6'>",
    "          <h1 className='text-2xl font-bold text-gray-800'>",
    "            Habit Tracker",
    "          </h1>",
    "          <p className='text-sm text-gray-600 mt-1'>",
    "            {user?.displayName}",
    "          </p>",
    "        </div>",
    "        ",
    "        <nav className='mt-6'>",
    "          {navItems.map(({ to, icon: Icon, label }) => (",
    "            <NavLink key={to} to={to}",
    "              className={({ isActive }) =>",
    "                `flex items-center px-6 py-3 text-gray-700",
    "                hover:bg-gray-100 ${isActive ?",
    "                'bg-gray-100 border-r-4 border-blue-500' : ''}`",
    "              }",
    "            >",
    "              <Icon className='w-5 h-5 mr-3' />",
    "              {label}",
    "            </NavLink>",
    "          ))}",
    "        </nav>",
    "",
    "        <button onClick={handleLogout}",
    "          className='flex items-center px-6 py-3 mt-auto",
    "                     text-red-600 hover:bg-red-50 w-full'",
    "        >",
    "          <LogOut className='w-5 h-5 mr-3' />",
    "          Logout",
    "        </button>",
    "      </aside>",
    "",
    "      <main className='flex-1 overflow-y-auto p-8'>",
    "        <Outlet />",
    "      </main>",
    "    </div>",
    "  )",
    "}"
], "JavaScript")

add_para(doc, 'Layout –≤–∏–∫–æ—Ä–∏—Å—Ç–æ–≤—É—î Flexbox –¥–ª—è –¥–≤–æ–∫–æ–ª–æ–Ω–∫–æ–≤–æ—ó —Å—Ç—Ä—É–∫—Ç—É—Ä–∏: —Ñ—ñ–∫—Å–æ–≤–∞–Ω–∞ –±—ñ—á–Ω–∞ –ø–∞–Ω–µ–ª—å (264px) —Ç–∞ –æ—Å–Ω–æ–≤–Ω–∏–π –∫–æ–Ω—Ç–µ–Ω—Ç (flex-1). NavLink –∞–≤—Ç–æ–º–∞—Ç–∏—á–Ω–æ –¥–æ–¥–∞—î –∫–ª–∞—Å –∞–∫—Ç–∏–≤–Ω–æ–≥–æ —Å—Ç–∞–Ω—É –¥–ª—è –ø–æ—Ç–æ—á–Ω–æ—ó —Å—Ç–æ—Ä—ñ–Ω–∫–∏. –ö–æ–º–ø–æ–Ω–µ–Ω—Ç Outlet –≤—ñ–¥ React Router —Ä–µ–Ω–¥–µ—Ä–∏—Ç—å –≤–∫–ª–∞–¥–µ–Ω—ñ –º–∞—Ä—à—Ä—É—Ç–∏. Tailwind CSS –∫–ª–∞—Å–∏ –∑–∞–±–µ–∑–ø–µ—á—É—é—Ç—å –∞–¥–∞–ø—Ç–∏–≤–Ω–∏–π –¥–∏–∑–∞–π–Ω.')

doc.add_paragraph()

add_explanation(doc,
    '–°—Ç–æ—Ä—ñ–Ω–∫–∞ Login.jsx - —Ñ–æ—Ä–º–∞ –∞–≤—Ç–æ—Ä–∏–∑–∞—Ü—ñ—ó:',
    'Login –∫–æ–º–ø–æ–Ω–µ–Ω—Ç —Ä–µ–∞–ª—ñ–∑—É—î —Å—Ç–æ—Ä—ñ–Ω–∫—É –≤—Ö–æ–¥—É –∑ –¥–≤–æ–∫–æ–ª–æ–Ω–∫–æ–≤–∏–º –º–∞–∫–µ—Ç–æ–º: –ª—ñ–≤–∞ —á–∞—Å—Ç–∏–Ω–∞ –∑ –æ–ø–∏—Å–æ–º —Ñ—É–Ω–∫—Ü—ñ–æ–Ω–∞–ª—É, –ø—Ä–∞–≤–∞ - –∑ —Ñ–æ—Ä–º–æ—é. –í–∏–∫–æ—Ä–∏—Å—Ç–æ–≤—É—î useState –¥–ª—è —É–ø—Ä–∞–≤–ª—ñ–Ω–Ω—è –ª–æ–∫–∞–ª—å–Ω–∏–º —Å—Ç–∞–Ω–æ–º —Ñ–æ—Ä–º–∏, react-hot-toast –¥–ª—è —Å–ø–æ–≤—ñ—â–µ–Ω—å. –ü—ñ—Å–ª—è —É—Å–ø—ñ—à–Ω–æ–≥–æ –≤—Ö–æ–¥—É –∑–±–µ—Ä—ñ–≥–∞—î —Ç–æ–∫–µ–Ω–∏ —É Zustand store —Ç–∞ –ø–µ—Ä–µ–Ω–∞–ø—Ä–∞–≤–ª—è—î –Ω–∞ dashboard.')

add_code_block(doc, [
    "import { useState } from 'react'",
    "import { Link, useNavigate } from 'react-router-dom'",
    "import { authAPI } from '../api/auth'",
    "import { useAuthStore } from '../store/authStore'",
    "import toast from 'react-hot-toast'",
    "import { Target } from 'lucide-react'",
    "",
    "export default function Login() {",
    "  const [email, setEmail] = useState('')",
    "  const [password, setPassword] = useState('')",
    "  const [loading, setLoading] = useState(false)",
    "  const navigate = useNavigate()",
    "  const { login } = useAuthStore()",
    "",
    "  const handleSubmit = async (e) => {",
    "    e.preventDefault()",
    "    setLoading(true)",
    "    try {",
    "      const data = await authAPI.login(email, password)",
    "      login(data.accessToken, data.refreshToken, data.user)",
    "      toast.success('Login successful!')",
    "      navigate('/dashboard')",
    "    } catch (error) {",
    "      toast.error(error.response?.data?.message || 'Login failed')",
    "    } finally {",
    "      setLoading(false)",
    "    }",
    "  }",
    "",
    "  return (",
    "    <div className='min-h-screen flex bg-gradient-to-br",
    "                    from-blue-500 via-purple-500 to-pink-500'>",
    "      <div className='flex-1 flex items-center justify-center p-8'>",
    "        <div className='max-w-md w-full bg-white rounded-2xl",
    "                        shadow-2xl p-8'>",
    "          <div className='text-center mb-8'>",
    "            <div className='inline-block p-3 bg-gradient-to-br",
    "                            from-blue-500 to-purple-500",
    "                            rounded-full mb-4'>",
    "              <Target className='w-8 h-8 text-white' />",
    "            </div>",
    "            <h2 className='text-3xl font-bold text-gray-800 mb-2'>",
    "              Welcome Back",
    "            </h2>",
    "          </div>",
    "          ",
    "          <form onSubmit={handleSubmit} className='space-y-5'>",
    "            <div>",
    "              <input type='email' value={email}",
    "                onChange={(e) => setEmail(e.target.value)}",
    "                className='w-full px-4 py-3 border-2",
    "                           border-gray-200 rounded-lg",
    "                           focus:outline-none focus:border-blue-500'",
    "                placeholder='you@example.com' required",
    "              />",
    "            </div>",
    "            <div>",
    "              <input type='password' value={password}",
    "                onChange={(e) => setPassword(e.target.value)}",
    "                className='w-full px-4 py-3 border-2",
    "                           border-gray-200 rounded-lg'",
    "                placeholder='‚Ä¢‚Ä¢‚Ä¢‚Ä¢‚Ä¢‚Ä¢‚Ä¢‚Ä¢' required",
    "              />",
    "            </div>",
    "            <button type='submit' disabled={loading}",
    "              className='w-full py-3 bg-gradient-to-r",
    "                         from-blue-600 to-purple-600",
    "                         text-white rounded-lg font-semibold'>",
    "              {loading ? 'Signing in...' : 'Sign In'}",
    "            </button>",
    "          </form>",
    "        </div>",
    "      </div>",
    "    </div>",
    "  )",
    "}"
], "JavaScript")

add_para(doc, '–§–æ—Ä–º–∞ –≤–∏–∫–æ—Ä–∏—Å—Ç–æ–≤—É—î controlled components pattern - –∑–Ω–∞—á–µ–Ω–Ω—è input\'—ñ–≤ –∑–±–µ—Ä—ñ–≥–∞—é—Ç—å—Å—è —É —Å—Ç–∞–Ω—ñ –∫–æ–º–ø–æ–Ω–µ–Ω—Ç–∞ —á–µ—Ä–µ–∑ useState. handleSubmit –∑–∞–ø–æ–±—ñ–≥–∞—î —Å—Ç–∞–Ω–¥–∞—Ä—Ç–Ω—ñ–π –ø–æ–≤–µ–¥—ñ–Ω—Ü—ñ —Ñ–æ—Ä–º–∏, –≤—Å—Ç–∞–Ω–æ–≤–ª—é—î loading —Å—Ç–∞–Ω, –≤–∏–∫–ª–∏–∫–∞—î API, —Ç–∞ –æ–±—Ä–æ–±–ª—è—î —É—Å–ø—ñ—Ö/–ø–æ–º–∏–ª–∫—É. –ì—Ä–∞–¥—ñ—î–Ω—Ç–Ω–∏–π —Ñ–æ–Ω —Å—Ç–≤–æ—Ä–µ–Ω–æ —á–µ—Ä–µ–∑ Tailwind CSS —É—Ç–∏–ª—ñ—Ç–∏.')

doc.add_page_break()

add_explanation(doc,
    '–°—Ç–æ—Ä—ñ–Ω–∫–∞ Dashboard.jsx - –≥–æ–ª–æ–≤–Ω–∞ –ø–∞–Ω–µ–ª—å –∫–æ—Ä–∏—Å—Ç—É–≤–∞—á–∞:',
    'Dashboard –≤—ñ–¥–æ–±—Ä–∞–∂–∞—î –æ–≥–ª—è–¥ —Ü—ñ–ª–µ–π –∫–æ—Ä–∏—Å—Ç—É–≤–∞—á–∞. –í–∏–∫–æ—Ä–∏—Å—Ç–æ–≤—É—î useEffect –¥–ª—è –∑–∞–≤–∞–Ω—Ç–∞–∂–µ–Ω–Ω—è –¥–∞–Ω–∏—Ö –ø—Ä–∏ –º–æ–Ω—Ç—É–≤–∞–Ω–Ω—ñ –∫–æ–º–ø–æ–Ω–µ–Ω—Ç–∞. –ü–æ–∫–∞–∑—É—î —Å—Ç–∞—Ç–∏—Å—Ç–∏–∫—É —É –≤–∏–≥–ª—è–¥—ñ –∫–∞—Ä—Ç–æ–∫ (–∑–∞–≥–∞–ª—å–Ω–∞ –∫—ñ–ª—å–∫—ñ—Å—Ç—å —Ü—ñ–ª–µ–π, –∞–∫—Ç–∏–≤–Ω—ñ, –∑–∞–≤–µ—Ä—à–µ–Ω—ñ) —Ç–∞ —Å–ø–∏—Å–æ–∫ —Ü—ñ–ª–µ–π —É –≤–∏–≥–ª—è–¥—ñ –∞–¥–∞–ø—Ç–∏–≤–Ω–æ—ó —Å—ñ—Ç–∫–∏.')

add_code_block(doc, [
    "import { useEffect, useState } from 'react'",
    "import { Link } from 'react-router-dom'",
    "import { goalsAPI } from '../api/goals'",
    "import { Plus, Target, TrendingUp } from 'lucide-react'",
    "import toast from 'react-hot-toast'",
    "",
    "export default function Dashboard() {",
    "  const [goals, setGoals] = useState([])",
    "  const [loading, setLoading] = useState(true)",
    "",
    "  useEffect(() => {",
    "    loadGoals()",
    "  }, [])",
    "",
    "  const loadGoals = async () => {",
    "    try {",
    "      const data = await goalsAPI.getAll()",
    "      setGoals(data)",
    "    } catch (error) {",
    "      toast.error('Failed to load goals')",
    "    } finally {",
    "      setLoading(false)",
    "    }",
    "  }",
    "",
    "  if (loading) {",
    "    return <div className='text-center py-12'>Loading...</div>",
    "  }",
    "",
    "  return (",
    "    <div>",
    "      <div className='mb-8'>",
    "        <h1 className='text-4xl font-bold bg-gradient-to-r",
    "                       from-blue-600 to-purple-600",
    "                       bg-clip-text text-transparent mb-2'>",
    "          Dashboard",
    "        </h1>",
    "      </div>",
    "",
    "      <div className='grid grid-cols-1 md:grid-cols-3 gap-6 mb-8'>",
    "        <div className='bg-gradient-to-br from-blue-500",
    "                        to-blue-600 rounded-2xl p-6",
    "                        text-white shadow-xl'>",
    "          <p className='text-blue-100 text-sm'>Total Goals</p>",
    "          <p className='text-3xl font-bold mt-1'>",
    "            {goals.length}",
    "          </p>",
    "        </div>",
    "        ",
    "        <div className='bg-gradient-to-br from-green-500",
    "                        to-green-600 rounded-2xl p-6",
    "                        text-white shadow-xl'>",
    "          <p className='text-green-100 text-sm'>Active Goals</p>",
    "          <p className='text-3xl font-bold mt-1'>",
    "            {goals.filter(g => g.status === 'ACTIVE').length}",
    "          </p>",
    "        </div>",
    "      </div>",
    "",
    "      <div className='grid grid-cols-1 md:grid-cols-2",
    "                      lg:grid-cols-3 gap-6'>",
    "        {goals.map((goal) => (",
    "          <Link key={goal.id} to={`/goals/${goal.id}`}",
    "            className='group bg-white rounded-2xl shadow-lg p-6",
    "                       hover:shadow-2xl transition-all",
    "                       transform hover:scale-105'>",
    "            <h3 className='text-lg font-bold text-gray-900",
    "                           group-hover:text-blue-600'>",
    "              {goal.title}",
    "            </h3>",
    "            <p className='text-gray-600 text-sm mt-2'>",
    "              {goal.description}",
    "            </p>",
    "          </Link>",
    "        ))}",
    "      </div>",
    "    </div>",
    "  )",
    "}"
], "JavaScript")

add_para(doc, 'useEffect –∑ –ø–æ—Ä–æ–∂–Ω—ñ–º –º–∞—Å–∏–≤–æ–º –∑–∞–ª–µ–∂–Ω–æ—Å—Ç–µ–π [] –≤–∏–∫–æ–Ω—É—î—Ç—å—Å—è –ª–∏—à–µ –ø—Ä–∏ –º–æ–Ω—Ç—É–≤–∞–Ω–Ω—ñ. loadGoals - async —Ñ—É–Ω–∫—Ü—ñ—è –¥–ª—è –∑–∞–≤–∞–Ω—Ç–∞–∂–µ–Ω–Ω—è –¥–∞–Ω–∏—Ö –∑ –æ–±—Ä–æ–±–∫–æ—é –ø–æ–º–∏–ª–æ–∫ —á–µ—Ä–µ–∑ try-catch. –°—Ç–∞—Ç–∏—Å—Ç–∏—á–Ω—ñ –∫–∞—Ä—Ç–∫–∏ –≤–∏–∫–æ—Ä–∏—Å—Ç–æ–≤—É—é—Ç—å filter –¥–ª—è –ø—ñ–¥—Ä–∞—Ö—É–Ω–∫—É —Ü—ñ–ª–µ–π –∑–∞ —Å—Ç–∞—Ç—É—Å–æ–º. Grid layout (grid-cols-1 md:grid-cols-3) –∑–∞–±–µ–∑–ø–µ—á—É—î –∞–¥–∞–ø—Ç–∏–≤–Ω—ñ—Å—Ç—å: 1 –∫–æ–ª–æ–Ω–∫–∞ –Ω–∞ –º–æ–±—ñ–ª—å–Ω–∏—Ö, 3 –Ω–∞ –¥–µ—Å–∫—Ç–æ–ø—ñ.')

doc.add_paragraph()

add_para(doc, '–¢–∞–∫–∏–º —á–∏–Ω–æ–º, frontend —á–∞—Å—Ç–∏–Ω–∞ —Ä–µ–∞–ª—ñ–∑–æ–≤–∞–Ω–∞ –Ω–∞ —Å—É—á–∞—Å–Ω–æ–º—É —Å—Ç–µ–∫—É React 18 + Vite –∑ –≤–∏–∫–æ—Ä–∏—Å—Ç–∞–Ω–Ω—è–º best practices: —Ñ—É–Ω–∫—Ü—ñ–æ–Ω–∞–ª—å–Ω—ñ –∫–æ–º–ø–æ–Ω–µ–Ω—Ç–∏ –∑ —Ö—É–∫–∞–º–∏, —Ü–µ–Ω—Ç—Ä–∞–ª—ñ–∑–æ–≤–∞–Ω–µ —É–ø—Ä–∞–≤–ª—ñ–Ω–Ω—è —Å—Ç–∞–Ω–æ–º —á–µ—Ä–µ–∑ Zustand, –º–æ–¥—É–ª—å–Ω–∞ –∞—Ä—Ö—ñ—Ç–µ–∫—Ç—É—Ä–∞ API –∫–ª—ñ—î–Ω—Ç—ñ–≤, –∞–≤—Ç–æ–º–∞—Ç–∏—á–Ω–µ –æ–Ω–æ–≤–ª–µ–Ω–Ω—è JWT —Ç–æ–∫–µ–Ω—ñ–≤, –∞–¥–∞–ø—Ç–∏–≤–Ω–∏–π –¥–∏–∑–∞–π–Ω —á–µ—Ä–µ–∑ Tailwind CSS.')

# –í–∏—Å–Ω–æ–≤–æ–∫ –¥–æ —Ä–æ–∑–¥—ñ–ª—É 4
doc.add_paragraph()
add_para(doc, '–£ –¥–∞–Ω–æ–º—É —Ä–æ–∑–¥—ñ–ª—ñ –¥–µ—Ç–∞–ª—å–Ω–æ —Ä–æ–∑–≥–ª—è–Ω—É—Ç–æ —Ä–µ–∞–ª—ñ–∑–∞—Ü—ñ—é –≤–µ–±-—Å–∏—Å—Ç–µ–º–∏ –¥–ª—è –≤—ñ–¥—Å—Ç–µ–∂–µ–Ω–Ω—è –∑–≤–∏—á–æ–∫. Backend –ø–æ–±—É–¥–æ–≤–∞–Ω–æ –Ω–∞ Spring Boot –∑ —Ç—Ä—å–æ—Ö—à–∞—Ä–æ–≤–æ—é –∞—Ä—Ö—ñ—Ç–µ–∫—Ç—É—Ä–æ—é (Controller-Service-Repository), –≤–∏–∫–æ—Ä–∏—Å—Ç–æ–≤—É—é—á–∏ MongoDB —è–∫ —Å—Ö–æ–≤–∏—â–µ –¥–∞–Ω–∏—Ö. Frontend —Ä–µ–∞–ª—ñ–∑–æ–≤–∞–Ω–æ –Ω–∞ React 18 –∑ –≤–∏–∫–æ—Ä–∏—Å—Ç–∞–Ω–Ω—è–º —Å—É—á–∞—Å–Ω–∏—Ö –ø—ñ–¥—Ö–æ–¥—ñ–≤: —Ñ—É–Ω–∫—Ü—ñ–æ–Ω–∞–ª—å–Ω—ñ –∫–æ–º–ø–æ–Ω–µ–Ω—Ç–∏, —Ö—É–∫–∏, Zustand –¥–ª—è state management, React Router –¥–ª—è –Ω–∞–≤—ñ–≥–∞—Ü—ñ—ó. –°–∏—Å—Ç–µ–º–∞ –∑–∞–±–µ–∑–ø–µ—á—É—î –±–µ–∑–ø–µ–∫—É —á–µ—Ä–µ–∑ JWT –∞–≤—Ç–µ–Ω—Ç–∏—Ñ—ñ–∫–∞—Ü—ñ—é, –∞–≤—Ç–æ–º–∞—Ç–∏—á–Ω–µ –æ–Ω–æ–≤–ª–µ–Ω–Ω—è —Ç–æ–∫–µ–Ω—ñ–≤, –≤–∞–ª—ñ–¥–∞—Ü—ñ—é –¥–æ—Å—Ç—É–ø—É –Ω–∞ —Ä—ñ–≤–Ω—ñ —Å–µ—Ä–≤—ñ—Å—ñ–≤. –ö–æ–¥ –Ω–∞–ø–∏—Å–∞–Ω–æ –∑ –¥–æ—Ç—Ä–∏–º–∞–Ω–Ω—è–º –ø—Ä–∏–Ω—Ü–∏–ø—ñ–≤ —á–∏—Å—Ç–æ–≥–æ –∫–æ–¥—É, SOLID —Ç–∞ DRY.')

# –ó–±–µ—Ä–µ–∂–µ–Ω–Ω—è –¥–æ–∫—É–º–µ–Ω—Ç—É
doc.save('/home/runner/work/cw_web/cw_web/–ö—É—Ä—Å–æ–≤–∞_–†–æ–±–æ—Ç–∞_HabitTracker.docx')
print("‚úÖ –î–æ–∫—É–º–µ–Ω—Ç —É—Å–ø—ñ—à–Ω–æ —Ä–æ–∑—à–∏—Ä–µ–Ω–æ!")
print("üìä –î–æ–¥–∞–Ω–æ –¥–µ—Ç–∞–ª—å–Ω—ñ –ø—Ä–∏–∫–ª–∞–¥–∏ –∫–æ–¥—É –∑ –ø–æ—è—Å–Ω–µ–Ω–Ω—è–º–∏:")
print("   - Entity –∫–ª–∞—Å–∏ (User, Goal, Group)")
print("   - Controller –∫–ª–∞—Å–∏ (Auth, Goal, Progress)")
print("   - Service –∫–ª–∞—Å–∏ (GoalService)")
print("   - React –∫–æ–º–ø–æ–Ω–µ–Ω—Ç–∏ (App, Layout, Login, Dashboard)")
print("   - API –∫–ª—ñ—î–Ω—Ç–∏ (client, goals)")
print("   - State management (Zustand)")
print("üìñ –ö–æ–∂–µ–Ω –ø—Ä–∏–∫–ª–∞–¥ –∫–æ–¥—É —Å—É–ø—Ä–æ–≤–æ–¥–∂—É—î—Ç—å—Å—è –¥–µ—Ç–∞–ª—å–Ω–∏–º –ø–æ—è—Å–Ω–µ–Ω–Ω—è–º")
