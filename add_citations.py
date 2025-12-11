#!/usr/bin/env python3
"""
Скрипт для додавання джерел (посилань) до теоретичних розділів курсової роботи.
Додає цитування до опису технологій та архітектурних паттернів.
"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn

# Визначаємо шлях до файлу
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
DOC_PATH = os.path.join(SCRIPT_DIR, 'Курсова_Робота_HabitTracker.docx')

def set_font(run, name='Times New Roman', size=14, bold=False, italic=False):
    """Встановлює шрифт для тексту."""
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), name)

def add_para(doc, text, bold=False, italic=False):
    """Додає параграф з форматуванням."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, bold=bold, italic=italic)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0.5)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    return p

# База джерел для технологій та архітектурних паттернів
REFERENCES = {
    # Backend технології
    'spring_boot': {
        'num': 1,
        'text': 'Spring Boot Reference Documentation. VMware, Inc., 2024. URL: https://docs.spring.io/spring-boot/docs/current/reference/html/ (дата звернення: 11.12.2024).'
    },
    'spring_security': {
        'num': 2,
        'text': 'Spring Security Reference. VMware, Inc., 2024. URL: https://docs.spring.io/spring-security/reference/ (дата звернення: 11.12.2024).'
    },
    'java': {
        'num': 3,
        'text': 'The Java™ Tutorials. Oracle Corporation, 2024. URL: https://docs.oracle.com/javase/tutorial/ (дата звернення: 11.12.2024).'
    },
    'jwt': {
        'num': 4,
        'text': 'Jones M., Bradley J., Sakimura N. JSON Web Token (JWT). RFC 7519, Internet Engineering Task Force, 2015. URL: https://datatracker.ietf.org/doc/html/rfc7519 (дата звернення: 11.12.2024).'
    },
    'mongodb': {
        'num': 5,
        'text': 'MongoDB Documentation. MongoDB, Inc., 2024. URL: https://www.mongodb.com/docs/ (дата звернення: 11.12.2024).'
    },
    'maven': {
        'num': 6,
        'text': 'Maven Documentation. Apache Software Foundation, 2024. URL: https://maven.apache.org/guides/ (дата звернення: 11.12.2024).'
    },
    'lombok': {
        'num': 7,
        'text': 'Project Lombok Documentation. The Project Lombok Authors, 2024. URL: https://projectlombok.org/features/ (дата звернення: 11.12.2024).'
    },
    
    # Frontend технології
    'react': {
        'num': 8,
        'text': 'React Documentation. Meta Platforms, Inc., 2024. URL: https://react.dev/learn (дата звернення: 11.12.2024).'
    },
    'vite': {
        'num': 9,
        'text': 'Vite Documentation. Evan You and Vite Contributors, 2024. URL: https://vitejs.dev/guide/ (дата звернення: 11.12.2024).'
    },
    'react_router': {
        'num': 10,
        'text': 'React Router Documentation. Remix Software Inc., 2024. URL: https://reactrouter.com/en/main (дата звернення: 11.12.2024).'
    },
    'zustand': {
        'num': 11,
        'text': 'Zustand Documentation. Poimandres, 2024. URL: https://docs.pmnd.rs/zustand/getting-started/introduction (дата звернення: 11.12.2024).'
    },
    'axios': {
        'num': 12,
        'text': 'Axios Documentation. Axios Contributors, 2024. URL: https://axios-http.com/docs/intro (дата звернення: 11.12.2024).'
    },
    'tailwind': {
        'num': 13,
        'text': 'Tailwind CSS Documentation. Tailwind Labs Inc., 2024. URL: https://tailwindcss.com/docs (дата звернення: 11.12.2024).'
    },
    
    # Безпека та алгоритми
    'bcrypt': {
        'num': 14,
        'text': 'Provos N., Mazières D. A Future-Adaptable Password Scheme. Proceedings of the USENIX Annual Technical Conference, 1999. URL: https://www.usenix.org/legacy/events/usenix99/provos/provos.pdf (дата звернення: 11.12.2024).'
    },
    'rest_api': {
        'num': 15,
        'text': 'Fielding R.T. Architectural Styles and the Design of Network-based Software Architectures. Doctoral dissertation, University of California, Irvine, 2000. URL: https://www.ics.uci.edu/~fielding/pubs/dissertation/top.htm (дата звернення: 11.12.2024).'
    },
    
    # Архітектурні паттерни
    'layered_architecture': {
        'num': 16,
        'text': 'Buschmann F., Meunier R., Rohnert H., Sommerlad P., Stal M. Pattern-Oriented Software Architecture, Volume 1: A System of Patterns. John Wiley & Sons, 1996. 476 p.'
    },
    'dependency_injection': {
        'num': 17,
        'text': 'Fowler M. Inversion of Control Containers and the Dependency Injection pattern, 2004. URL: https://martinfowler.com/articles/injection.html (дата звернення: 11.12.2024).'
    },
    'repository_pattern': {
        'num': 18,
        'text': 'Fowler M. Patterns of Enterprise Application Architecture. Addison-Wesley Professional, 2002. 560 p.'
    },
    'builder_pattern': {
        'num': 19,
        'text': 'Gamma E., Helm R., Johnson R., Vlissides J. Design Patterns: Elements of Reusable Object-Oriented Software. Addison-Wesley Professional, 1994. 416 p.'
    },
    'component_architecture': {
        'num': 20,
        'text': 'Abramov D., Clark A. Thinking in React. React Documentation, Meta Platforms, Inc., 2024. URL: https://react.dev/learn/thinking-in-react (дата звернення: 11.12.2024).'
    },
}

# Мапінг технологій до посилань
TECH_CITATIONS = {
    'Spring Boot 3.2.0': '[1]',
    'Java 17': '[3]',
    'Spring Security з JWT': '[2, 4]',
    'MongoDB': '[5]',
    'Maven': '[6]',
    'Lombok': '[7]',
    'React 18': '[8]',
    'Vite': '[9]',
    'React Router v6': '[10]',
    'Zustand': '[11]',
    'Axios': '[12]',
    'Tailwind CSS': '[13]',
    'BCrypt': '[14]',
    'REST API': '[15]',
    'RESTful API': '[15]',
}

PATTERN_CITATIONS = {
    'Багатошарова архітектура': '[16]',
    'Layered Architecture': '[16]',
    'Dependency Injection': '[17]',
    'Repository Pattern': '[18]',
    'Builder Pattern': '[19]',
    'Component-Based Architecture': '[20]',
}

def add_citation_to_text(text, citations_map):
    """Додає посилання до тексту на основі знайдених технологій/паттернів."""
    modified_text = text
    added_citations = set()
    
    # Сортуємо за довжиною (довші спочатку) щоб уникнути часткових замін
    sorted_items = sorted(citations_map.items(), key=lambda x: len(x[0]), reverse=True)
    
    for term, citation in sorted_items:
        if term in modified_text and citation not in added_citations:
            # Додаємо посилання після терміна
            modified_text = modified_text.replace(term, f"{term} {citation}", 1)
            added_citations.add(citation)
    
    return modified_text

def process_document():
    """Обробляє документ та додає посилання."""
    print("Завантаження документу...")
    doc = Document(DOC_PATH)
    
    print("Додавання посилань до тексту...")
    
    # Обробляємо параграфи з технологіями (приблизно 125-145)
    # та архітектурними паттернами (приблизно 656-668, 773-785)
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        
        # Пропускаємо заголовки та порожні параграфи
        if not text or len(text) < 50:
            continue
        
        # Перевіряємо чи це параграф з технологіями
        if any(tech in text for tech in TECH_CITATIONS.keys()):
            new_text = add_citation_to_text(text, TECH_CITATIONS)
            if new_text != text:
                # Замінюємо текст параграфа
                para.clear()
                run = para.add_run(new_text)
                set_font(run)
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                para.paragraph_format.first_line_indent = Inches(0.5)
                para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                print(f"  Додано посилання в параграф {i}")
        
        # Перевіряємо чи це параграф з архітектурними паттернами
        if any(pattern in text for pattern in PATTERN_CITATIONS.keys()):
            new_text = add_citation_to_text(text, PATTERN_CITATIONS)
            if new_text != text:
                para.clear()
                run = para.add_run(new_text)
                set_font(run)
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                para.paragraph_format.first_line_indent = Inches(0.5)
                para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                print(f"  Додано посилання в параграф {i}")
    
    print("\nДодавання розділу 'Перелік посилань'...")
    
    # Додаємо розділ з переліком посилань в кінець документу
    doc.add_page_break()
    
    # Заголовок
    heading = doc.add_heading(level=1)
    run = heading.add_run('ПЕРЕЛІК ПОСИЛАНЬ')
    set_font(run, size=16, bold=True)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Сортуємо посилання за номером
    sorted_refs = sorted(REFERENCES.values(), key=lambda x: x['num'])
    
    # Додаємо кожне посилання
    for ref in sorted_refs:
        p = doc.add_paragraph()
        run = p.add_run(f"{ref['num']}. {ref['text']}")
        set_font(run, size=14)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)  # Висячий відступ
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.space_after = Pt(6)
    
    print(f"\nДодано {len(sorted_refs)} джерел до переліку посилань")
    
    # Зберігаємо документ
    print("\nЗбереження документу...")
    doc.save(DOC_PATH)
    
    print("✅ Документ успішно оновлено!")
    print(f"   Посилання додано до технологій та архітектурних паттернів")
    print(f"   Створено розділ 'Перелік посилань' з {len(sorted_refs)} джерелами")
    print(f"\n📄 Оновлений файл: {DOC_PATH}")

if __name__ == '__main__':
    try:
        process_document()
    except Exception as e:
        print(f"❌ Помилка: {e}")
        import traceback
        traceback.print_exc()
        exit(1)
